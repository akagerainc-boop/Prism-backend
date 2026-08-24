"""Word (.docx) reconstruction via python-docx.

Preserves the visual hierarchy the layout model recovered: titles and headings
become real Word heading styles (so the navigation pane and any generated TOC
work), tables become real Word tables with merged cells, and figures are
embedded rather than dropped.
"""

from __future__ import annotations

import io

from ..document_model import table_rows
from ..logging_config import get_logger
from ..schemas import DocElement, StructuredDocument
from .images import element_image_bytes

log = get_logger(__name__)


class ExportUnavailable(RuntimeError):
    """A required export library is not installed."""


def _import_docx():
    try:
        import docx  # type: ignore[import-not-found]
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore

        return docx, WD_ALIGN_PARAGRAPH, Inches, Pt
    except Exception as exc:
        raise ExportUnavailable(
            "python-docx is not installed. Run `pip install python-docx`."
        ) from exc


def _add_table(doc, element: DocElement, Inches) -> None:
    table_data = element.table
    if table_data is None:
        return

    rows = table_rows(table_data)
    if not rows:
        if element.html:
            doc.add_paragraph(element.html)
        return

    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            text = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.text = text or ""

    # Re-apply the merges the layout model detected.
    for cell_spec in table_data.cells:
        if cell_spec.rowSpan <= 1 and cell_spec.colSpan <= 1:
            continue
        top = cell_spec.row
        left = cell_spec.col
        bottom = min(top + cell_spec.rowSpan - 1, len(rows) - 1)
        right = min(left + cell_spec.colSpan - 1, column_count - 1)
        if bottom <= top and right <= left:
            continue
        try:
            origin = table.cell(top, left)
            target = table.cell(bottom, right)
            merged = origin.merge(target)
            merged.text = cell_spec.text or ""
        except (IndexError, ValueError) as exc:  # pragma: no cover
            log.debug("Skipped a table merge (%s)", exc)

    # Bold the header row when one was identified.
    if any(c.isHeader and c.row == 0 for c in table_data.cells):
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True


def _add_image(doc, element: DocElement, document: StructuredDocument, Inches) -> bool:
    data = element_image_bytes(element, document)
    if not data:
        return False
    try:
        doc.add_picture(io.BytesIO(data), width=Inches(5.5))
        return True
    except Exception as exc:  # pragma: no cover - unsupported image payload
        log.debug("Could not embed image for element %s: %s", element.id, exc)
        return False


def export_docx(document: StructuredDocument, **_: object) -> bytes:
    """Render the structured document to .docx bytes."""
    docx, WD_ALIGN_PARAGRAPH, Inches, Pt = _import_docx()

    doc = docx.Document()

    for page_index, page in enumerate(document.pages):
        if page_index > 0:
            doc.add_page_break()

        for element in sorted(page.elements, key=lambda e: e.readingOrder):
            text = (element.text or "").strip()

            element_type = element.type
            if element_type == "title":
                if text:
                    doc.add_heading(text, level=0)
            elif element_type == "heading":
                if text:
                    doc.add_heading(text, level=min(max(element.level or 1, 1), 9))
            elif element_type == "table":
                _add_table(doc, element, Inches)
            elif element_type in ("figure", "chart"):
                if not _add_image(doc, element, document, Inches) and text:
                    doc.add_paragraph(text)
            elif element_type == "caption":
                if text:
                    paragraph = doc.add_paragraph(text)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.italic = True
                        run.font.size = Pt(9)
            elif element_type == "list":
                for line in text.splitlines():
                    if line.strip():
                        doc.add_paragraph(line.strip(), style="List Bullet")
            elif element_type == "formula":
                if text:
                    paragraph = doc.add_paragraph(text)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.italic = True
            elif element_type in ("footnote", "reference"):
                if text:
                    paragraph = doc.add_paragraph(text)
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
            elif element_type in ("header", "footer"):
                continue  # page furniture, not body content
            else:
                if text:
                    doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
